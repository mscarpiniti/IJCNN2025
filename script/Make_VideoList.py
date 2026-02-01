# -*- coding: utf-8 -*-
"""
This script is used to generate a list of the video presentations in the conference
program, from the report of accepted papers containing the full information.
It is used with two excel files:
- a file downloaded from Microsoft CMT system containing all the accepted papers,
- a file containing the IDs of all papers presented as a video

Created on Mon May 26 14:51:34 2025

@author: Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
"""

import pandas as pd
import IJCNN_utils as ut
from docx import Document


# Set file names
path = './data/'
dest = './report/'
paper_file = path + 'All_Accepted_Papers.xlsx'
video_file = path + 'Video_Presentation.xlsx'

save_file  = dest + 'Video_Program.docx'


# Load paper data
papers = ut.load_submissions(paper_file)
papers.set_index('ID', inplace=True)


# Load remote papers
videos = pd.read_excel(video_file)
video_list = list(videos['ID'])


# Create a new document
document = Document()


# Add the document header
document.add_heading('IJCNN 2025 Program -- Video Presentations')

p = document.add_paragraph('')
p.add_run('').italic = True


# Make the video list
ut.make_video_list(video_list, papers, document)


# Save the docx document
document.save(save_file)
