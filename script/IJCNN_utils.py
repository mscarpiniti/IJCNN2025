# -*- coding: utf-8 -*-
"""
This file contains some useful functions used in the managing of the IJCNN 2025
conference

Created on Tue Mar 11 22:22:01 2025

@author: Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from email_validator import validate_email, EmailNotValidError
from matplotlib.backends.backend_pdf import PdfPages
import pymupdf as pdf



# %% #### GENERAL FUNCTIONS ####

# Function used for saving all figures in a single pdf file
def save_all_figs(filename, figs=None, dpi='figure'):
    """
    Function used for saving all opened figures in a single pdf file.

    Parameters
    ----------
    filename : string
        Name of the file where saving all figures.
    figs : list, optional
        List of figures to save in the file. If None all figurea are saved.
        The default is None.
    dpi : int, optional
        he resolution in dots per inch. If 'figure', use the figure's dpi value.
        The default is 'figure'.

    Returns
    -------
    None.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    pp = PdfPages(filename)
    if figs is None:
        figs = [plt.figure(n) for n in plt.get_fignums()]
    for fig in figs:
        fig.savefig(pp, format='pdf', dpi=dpi)
    pp.close()




# %% #### REVIEWERS FUNCTIONS ####

# Function for check if an email address is valid
def check_email(email):
    """
    Function for check if an email address is valid or not.

    Parameters
    ----------
    email : string
        Email address to be checked.

    Returns
    -------
    bool
        If the email address is valid or not.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    try:
        valid = validate_email(email)
        return True
    except EmailNotValidError as e:
        return False



# Function to test a list of emails
def test_emails(em_list):
    """
    Function to check a list of emails

    Parameters
    ----------
    em_list : list
        List of email addresses to be checked.

    Returns
    -------
    r : List of booleans
        List of the check results, if the corresponding address is valid or not.
    res : List of integers
          List of the indices of r, corresponding to addresses not valid.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    r = []
    for m in em_list:
        rr = check_email(m)
        r.append(rr)

    res = [i for i, val in enumerate(r) if not val]

    return r, res



# Function to split an author in: name, middle initial, and surname
def split_name(x):
    """
    Function to split an author in: name, middle initial, and surname

    Parameters
    ----------
    x : string
        a string containing a full name.

    Returns
    -------
    n : string
        a string containing the first name.
    m : string
        a string containing the middle initial.
    s : string
        a string containing the last name.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    x = x.split()
    n = x[0].title()
    s = x[-1].title()
    if len(x) > 2:
        m = x[1][0].upper() + '.'
    else:
        m = ''

    return n, m, s



# Function to split the first name in first name and middle initial
def split_firstName(df):
    """
    Function used to split a first name into its first name and middle initial.

    Parameters
    ----------
    df : DataFrame
        DataFrame containing the Reviewers' information.

    Returns
    -------
    first_name : List
        A list containing all the first names.
    mid_name : List
        A list containing all the middle initials, if available, or an empty
        entry otherwise.

    """
    first_name = []
    mid_name = []

    for i in range(len(df)):
        x = df.iloc[i]['First Name']
        x = x.split()
        first_name.append(x[0].title())
        if len(x) > 1:
            mid_name.append(x[1][0].upper() + '.')
        else:
            mid_name.append('')

    return first_name, mid_name



# Function for menaging the middle name initial
def get_initial(df):
    """
    Function for menaging the middle name initial: it converts the full middle
    name of a DataFrame into the corresponding middle initials.

    Parameters
    ----------
    df : DataFrame
        DataFrame containing users' data. In particular it must contain a column
        named as 'Middle Initial'.

    Returns
    -------
    df : DataFrame
        The same input DataFrame where the 'Middle Initial' column contains the
        middle initials and no longer the full middle name.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    for k in range(len(df)):
        x = df.iloc[k]['Middle Initial']
        if not pd.isna(x):
            df.iloc[k]['Middle Initial'] = x[0].upper() + '.'

    return df



# Function to get the Organization name
def get_organization(authors):
    """
    Function to get the main Organization name.

    Parameters
    ----------
    authors : string
        string containing the list of all authors including the affiliations.
        This string is extracted from the column "Authors" of the file downloaded
        from CMT. The main organization (related to the main contact author) is
        denoted with a * symbol.

    Returns
    -------
    o : string
        the Organization name.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    x = authors.split('*')
    z = x[0].split('(')

    o = z[-1][:-1]

    return o



# Function to extract reviewers and area chairs
def extract_users(df):
    """
    Function to extract the list of reviewers or area chairs.

    Parameters
    ----------
    df : DataFrame
        DataFrame containing users' data. In particular it must contain columns
        named as 'First Name', 'Last Name', and'Middle Initial'.

    Returns
    -------
    U_ser : Series
        Series containing the full name of  the users.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    N = len(df)
    U = []

    for i in range(N):
        if str(df.iloc[i]['Middle Initial (optional)']) == 'nan':
            name = df.iloc[i]['First Name'].strip() + ' ' + df.iloc[i]['Last Name'].strip()
        else:
            name = df.iloc[i]['First Name'].strip() + ' ' + df.iloc[i]['Middle Initial (optional)'] + ' ' + df.iloc[i]['Last Name'].strip()

        U.append(name.title())

    U_ser = pd.Series(U)

    return U_ser



# %% #### REVIEWS FUNCTIONS ####

# Function for loading all reviews
def load_reviews(filename='Reviews.xlsx', header=0, sheet=0):
    """
    Function used for loading all reviews extracted from CMT.

    Parameters
    ----------
    filename : string, optional
        The name of the file to open. The default is 'Reviews.xlsx'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).

    Returns
    -------
    data : DataFrame
        a DataFrame containing all the useful information. This is limited only
        to a subset of 10 columns of the whole data.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    columns_select = ['Paper ID', 'Paper Title', 'Reviewer Email',
    'Q1 (Reviewer’s confidence - Value)', 'Q2 (Relevance to IJCNN - Value)',
    'Q3 (Technical quality - Value)', 'Q4 (Novelty - Value)',
    'Q5 (Quality of presentation - Value)', 'Q6 (Award quality? - Value)',
    'Q9 (Overall recommendation - Value)']

    column_names = ['ID', 'Title', 'Email', 'Rc', 'C1', 'C2', 'C3', 'C4', 'AQ', 'R']

    data = pd.read_excel(filename, header=header, sheet_name=sheet)
    data = data[columns_select]
    data.columns = column_names

    nan = data.isnull().sum().sum()
    if nan:
        print('There are %d missing values' % nan)

    return data



# Function for loading all reviews
def load_metareviews(filename='MetaReviews.xlsx', header=0, sheet=0):
    """
    Function used for loading all meta-reviews extracted from CMT.

    Parameters
    ----------
    filename : string, optional
        The name of the file to open. The default is 'MetaReviews.xlsx'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).

    Returns
    -------
    data : DataFrame
        a DataFrame containing all the useful information. This is limited only
        to a subset of 6 columns of the whole data.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """

    columns_select = ["Paper ID", "Paper Title", "Meta-Reviewer Email",
    "Q2 (Overall recommendation  - Value)", "Q4 (Paper formatting needs adjustment?)",
    "Q5 (Was Authors' anonymity ensured?  (If No, please explain in Comments to Chairs))"]

    column_names = ['ID', 'Title', 'Email', 'R', 'F', 'Anonimity']

    data = pd.read_excel(filename, header=header, sheet_name=sheet)
    data = data[columns_select]
    data.columns = column_names

    nan = data.isnull().sum().sum()
    if nan:
        print('There are %d missing values' % nan)

    return data



# Funtion for loading all the submissions
def load_submissions(filename='Papers.xlsx', header=0, sheet=0, other_col=None, other_names=None):
    """
    Funtion for loading all the submissions

    Parameters
    ----------
    filename : string, optional
        The name of the file to open. The default is 'MetaReviews.xlsx'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).
    other_col : list, optional
        other columns to extract except the 12 considered by default.
        The default is None.
    other_names : list, optional
        names of the other columns to extract except the 12 considered by default.
        The default is None.

    Returns
    -------
    data : DataFrame
        a DataFrame containing all the loaded information. By defaut, this is
        limited only to a subset of 12 columns of the whole data.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """
    columns_select = ["Paper ID", "Paper Title", "Abstract",
                      "Primary Contact Author Name", "Primary Contact Author Email",
                      "Authors", "Author Names", "Author Emails", "Track Name",
                      "Primary Subject Area", "Secondary Subject Areas", "Status"]

    column_names = ["ID", "Title", "Abstract", "Primary Contact",
                    "Primary Contact Email", "Authors", "Author Names",
                    "Author Emails", "Track", "Primary Subject Area",
                    "Secondary Subject Areas", "Status"]

    data = pd.read_excel(filename, header=header, sheet_name=sheet)

    if other_col is not None:
        columns_select = columns_select + other_col
        if other_names is not None:
            column_names = column_names + other_names
        else:
            column_names = column_names + other_col

    data = data[columns_select]
    data.columns = column_names

    nan = data.isnull().sum().sum()
    if nan:
        print('There are %d missing values' % nan)

    return data



# Function for loading all reviewer ratings
def load_reviewratings(filename='ReviewRatings.txt'):
    """
    Function for loading all the reviewer ratings extracted from CMT.

    Parameters
    ----------
    filename : string, optional
        tha name of the file to be loaded. The default is 'ReviewRatings.txt'.

    Returns
    -------
    data : DataFrame
        a DataFrame containing all the loaded information.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """
    column_names = ['Reviewer', 'Rating', 'ID', 'Meta-Reviewer']

    data = pd.read_csv(filename, sep='\t')
    data.columns = column_names

    nan = data.isnull().sum().sum()
    if nan:
        print('There are %d missing values' % nan)

    return data



# Function for loading conference Users
def load_users(filename='Users.txt'):
    """
    Function for loading all the conference Users.

    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Users.txt'.

    Returns
    -------
    data : DataFrame
        a DataFrame containing all the loaded information. By defaut, this is
        limited only to a subset of 7 columns of the whole data.

    Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
    """
    columns_select = ['# First Name', 'Middle Initial (optional)', 'Last Name',
                      'Email', 'Organization', 'Country', 'Roles']

    column_names = ['First Name', 'Middle Initial', 'Last Name',
                      'Email', 'Organization', 'Country', 'Roles']

    data = pd.read_csv(filename, sep='\t')
    data = data[columns_select]
    data.columns = column_names

    nan = data.isnull().sum().sum()
    if nan:
        print('There are %d missing values' % nan)

    return data



# Function for loading conference Reviewers/Meta-Reviewers
def load_reviewers(filename='Reviewers.txt', sheet=0):
    """
    Function for loading all the conference Reviewers/Meta-Reviewers.

    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Reviewers.txt'.
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).

    Returns
    -------
    data : DataFrame
        a DataFrame containing all the loaded information.

    """
    ext = filename.split('.')[-1]
    if ext == 'txt':
        data = pd.read_csv(filename, sep='\t')
    elif ext == 'xlsx':
        data = pd.read_excel(filename, sheet_name=sheet)
    else:
        print('Error: file type unsupported!')

    nan = data.isnull().sum().sum()
    if nan:
        print('There are %d missing values' % nan)

    return data




# Convert a CMT csv file to an xlsx one
def csv2xlsx(filename='Reviews.csv', header=2):
    """
    Function used to convert a CMT csv file into an Excel version.

    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Reviews.csv'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).

    Returns
    -------
    None.

    """
    df = pd.read_csv(filename, header)

    savename = filename.split('.')[0] + '.xlsx'
    df.to_excel(savename)



# Convert a CMT xls file to an xlsx one
def xls2xlsx(filename='Reviews.xls', header=2):
    """
    Function used to convert a CMT old xls file into the new xlsx Excel version.


    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Reviews.xls'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).

    Returns
    -------
    None.

    """
    df = pd.read_excel(filename, header, engine='xlrd')

    savename = filename.split('.')[0] + '.xlsx'
    df.to_excel(savename)



# Function to retrieve paper status
def get_status(filename='Papers.xlsx', header=0, sheet=0):
    """
    Function used to obtain the paper status for each ID.

    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Papers.xlsx'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).

    Returns
    -------
    data : DataFrame
        A DataFrame containing the paper ID (first column) and related Status (second column).

    """
    df = pd.read_excel(filename, header=header, sheet_name=sheet)
    df = df.rename(columns={'Paper ID': 'ID'})
    data = df[['ID', 'Status']]

    return data



# Function to retrieve paper authors
def get_authors(filename='Papers.xlsx', header=0, sheet=0):
    """
    Function used to obtain the paper authors for each ID.

    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Papers.xlsx'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).

    Returns
    -------
    data : DataFrame
        A DataFrame containing the paper ID (first column) and related Authors (second column).

    """
    df = pd.read_excel(filename, header=header, sheet_name=sheet)
    df = df.rename(columns={'Paper ID': 'ID'})
    data = df[['ID', 'Author Names']]

    return data



# Function to retrieve paper authors and status
def get_authors_status(filename='Papers.xlsx', header=0, sheet=0):
    """
    Function used to obtain the paper authors and related status for each ID.

    Parameters
    ----------
    filename : string, optional
        the name of the file to be loaded. The default is 'Papers.xlsx'.
    header : integer, optional
        The number of the row where finding the header. The default is 0.
        Usually, CMT file start from the third row (header=2).
    sheet : integer or string, optional
        the number or name of the file sheet to read. The default is 0 (the first one).

    Returns
    -------
    data : DataFrame
        A DataFrame containing the paper ID (first column), the related Authors (second column),
        and the corresponding status (third column).

    """
    df = pd.read_excel(filename, header=header, sheet_name=sheet)
    df = df.rename(columns={'Paper ID': 'ID'})
    data = df[['ID', 'Author Names', 'Status']]
    data = data.rename(columns={'Author Names': 'Authors'})

    return data




# Function used to de-quantize per-paper review scores
def dequantize_score(s, lamb=25, q=0.5):
    """
    Function used for review score de-quantization.

    Parameters
    ----------
    s : list
        List of review scores.
    lamb : integer, optional
        Lambda tradeoff hyper-parameter. The default is 25.
    q : float, optional
        Width of each quantized interval. The default is 0.5.

    Returns
    -------
    y : List
        List of de-quantized review scores.

    """
    y = []
    N = len(s)
    for i in range(N):
        den = N*(1 + lamb)
        y_tilde_1 = (1 + N*lamb)/den * s.iloc[i]
        y_tilde_2 = (s.sum() - s.iloc[i])/den
        y_tilde = y_tilde_1 + y_tilde_2

        y.append(min(max(y_tilde, s.iloc[i]-q/2), s.iloc[i]+q/2))

    return y



# Function used to de-quantize all review scores
def dequantize(data, lamb=25, q=0.5):
    """
    Function used for the de-quantization of all received review reports.

    Parameters
    ----------
    data : DataFrame
        DataFrame containing both the paper ID and related received score (Sr).
    lamb : integer, optional
        Lambda tradeoff hyper-parameter. The default is 25.
    q : float, optional
        Width of each quantized interval. The default is 0.5.

    Returns
    -------
    data : DataFrame
        DataFrame containing the paper ID, the received score (Sr) and a new
        column with the de-quantized score (Sr_dq).

    """
    data_Sr_dq = data.groupby(by=['ID'])['Sr'].apply(dequantize_score, lamb, q)

    Sr_dq = []
    for i in range(len(data_Sr_dq)):
        for j in range(len(data_Sr_dq.iloc[i])):
            Sr_dq.append(data_Sr_dq.iloc[i][j])

    Sr_dq = pd.Series(Sr_dq)
    Sr_dq.name = 'Sr_dq'

    data = data.join(Sr_dq)

    return data



# Function for computing a weighted mean
def weighted_mean(x, w):
    """
    Function used to evaluate a weighted mean between two vectors.

    Parameters
    ----------
    x : Array
        Array of values.
    w : Array
        Array of weights.

    Returns
    -------
    z : Float
        Weighted mean.

    """
    z = np.sum(w*x)/np.sum(w)

    return z



# Function for computing a weighted mean (for DataFrame)
def weighted_avg(df, values, weights):
    """
    Function used to evaluate a weighted mean between weights and values which
    are passed as columns of a DataFrame.

    Parameters
    ----------
    df : DataFrame
        Dataframe containing weights and values.
    values : string
        Name of the DataFrame column containing the values.
    weights : string
        Name of the DataFrame column containing the weights.

    Returns
    -------
    Series
        Series containing the evaluated weighted mean.

    """
    d = df[values]
    w = df[weights]

    m = (d * w).sum() / w.sum()

    return m



# Function for computing a harmonic mean (for DataFrame)
def harmonic_avg(df, val1, val2):
    """
    Function used to evaluate the harmonic mean between two sets of data which
    are passed as columns of a DataFrame.

    Parameters
    ----------
    df : DataFrame
        Dataframe containing weights and values.
    val1 : string
        Name of the DataFrame column containing the first set of data.
    val2 : string
        Name of the DataFrame column containing the second set of data.

    Returns
    -------
    Series
        Series containing the evaluated harmonic mean.

    """
    a = df[val1]
    b = df[val2]

    m = (2 * a * b) / (a + b)

    return m



# Function for computing a power mean (for DataFrame)
def power_avg(df, val1, val2, p=2):
    """
    Function used to evaluate the power mean (with power p) between two sets of
    data which are passed as columns of a DataFrame.

    Parameters
    ----------
    df : DataFrame
        Dataframe containing weights and values.
    val1 : string
        Name of the DataFrame column containing the first set of data.
    val2 : string
        Name of the DataFrame column containing the second set of data.
    p : integer
        Value of the power used in the mean evaluation.

    Returns
    -------
    Series
        Series containing the evaluated power mean.

    """
    a = df[val1]
    b = df[val2]

    m = (0.5*(a**p + b**p))**(1/p)

    return m



# Function for normalizing data
def normalize(s):
    """
    Function used for the data normalization: i.e., mapping all values inside
    the [0, 1] interval.

    Parameters
    ----------
    s : Array
        Input array of data.

    Returns
    -------
    n : Array
        Normalized array.

    """
    n = (s - s.min())/(s.max() - s.min())

    return n



# Function for calibrating the reviews
def calibration(reviews, column='Sr_dq', number_accepts=2000, who='Reviewer'):
    """
    Function used to implement the experimental claibration of reviewers, as suggested in:
        - C. Cortes and N. D. Lawrence, “Inconsistency in conference peer review:
            Revisiting the 2014 NeurIPS experiment,” arXiv:2109.09774, 2021.
        - N. B. Shah, B. Tabibian, K. Muandet, I. Guyon, and U. V. Luxburg,
            “Design and analysis of the NIPS 2016 review process,” The Journal of
            Machine Learning Research, vol. 19, no. 1, pp. 1913–1946, 2018.
        - https://inverseprobability.com/2014/08/02/reviewer-calibration-for-nips
        - https://github.com/lawrennd/neurips2014/blob/master/notebooks/Reviewer%20Calibration.ipynb

    The function also generates some plots on calibration statistics and acceptance
    probability.

    Parameters
    ----------
    reviews : DataFrame
        DataFrame from CMT containing all the reviews.
    column : string, optional
        Name of the column containing the score to be calibrated. The default is 'Sr_dq'.
    number_accepts : integer, optional
        Number of paper to be accepted. The default is 2000.
    who : string, optional
        A string denoting Reviewer/Meta-reviewer, for printing the correct users.
        The default is 'Reviewer'.

    Returns
    -------
    reviews : DataFrame
        DataFrame containing the calibrated scores with an estimated acceptance probability.
    revs : DataFrame
        DataFrame containing the ordered calibrated scores with an estimated
        acceptance probability in descending order, along with the estimated
        mean and std of reviewers' biases.

    """

    import GPy

    # Evaluate the mean
    mu = reviews[column].mean()
    print("Mean value, mu =", mu)

    # Data preparation
    X1 = pd.get_dummies(reviews.ID)
    X1 = X1[sorted(X1.columns, key=int)]
    X2 = pd.get_dummies(reviews.Email)
    X2 = X2[sorted(X2.columns, key=str.lower)]
    X  = X1.join(X2)
    y  = reviews[column] - mu

    # Constructing the model in GPy
    kern1 = GPy.kern.Linear(input_dim=len(X1.columns), active_dims=np.arange(len(X1.columns)))
    kern1.name = 'K_f'
    kern2 = GPy.kern.Linear(input_dim=len(X2.columns), active_dims=np.arange(len(X1.columns), len(X.columns)))
    kern2.name = 'K_b'

    model = GPy.models.GPRegression(X, y.to_numpy()[:, np.newaxis], kern1+kern2)
    model.optimize()

    print(model)
    print(model.log_likelihood())

    # Constructing model without GPy
    # set parameter values to ML solutions given by GPy.
    alpha_f = model.sum.K_f.variances
    alpha_b = model.sum.K_b.variances/alpha_f
    sigma2  = model.Gaussian_noise.variance/alpha_f

    K_f = np.dot(X1, X1.T)
    K_b = alpha_b*np.dot(X2, X2.T)
    K = K_f + K_b + sigma2*np.eye(X2.shape[0])
    Kinv, L, Li, logdet = GPy.util.linalg.pdinv(K)
    alpha = np.dot(Kinv, y)
    yTKinvy = np.dot(y, alpha)
    alpha_f = yTKinvy/len(y)

    ll = 0.5*len(y)*np.log(2*np.pi*alpha_f) + 0.5*logdet + 0.5*yTKinvy/alpha_f
    print("Negative log likelihood: ", ll)

    # Compute mean and covariance of quality scores
    K_s = K_f + np.eye(K_f.shape[0])*sigma2
    s = pd.Series(np.dot(K_s, alpha) + mu, index=X1.index)
    covs = alpha_f*(K_s - np.dot(K_s, np.dot(Kinv, K_s)))


    cal_name = column + '_Calibrated'
    s.name = cal_name
    reviews = reviews.join(s)


    # Monte Carlo Simulations for Probability of Acceptance
    samples = 1000
    score = np.random.multivariate_normal(mean=s, cov=covs, size=samples).T
    # Use X1 which maps papers to paper/reviewer pairings to get the average score for each paper.
    paper_score = pd.DataFrame(np.dot(np.diag(1./X1.sum(0)), np.dot(X1.T, score)), index=X1.columns)

    prob_accept = ((paper_score>paper_score.quantile(1-(float(number_accepts)/paper_score.shape[0]))).sum(1)/1000)
    prob_name = column + '_AcceptProbability'
    prob_accept.name = prob_name

    lower = 0.1
    upper = 0.9
    grey_area = ((prob_accept>lower) & (prob_accept<upper))
    print('Number of papers in grey area:', grey_area.sum())

    fig, ax = plt.subplots()
    print('Expected Papers Accepted:', prob_accept.sum())
    _ = prob_accept.hist(bins=40, ax=ax)
    _ = ax.set_title('Probability of Acceptance', fontsize=18)


    # Some sanity histograms
    fig, ax = plt.subplots()
    s.hist(bins=100, ax=ax)
    _ = ax.set_title('Calibrated ' + who + ' Scores', fontsize=18)

    # Adjustments to Reviewer Scores
    # Compute mean and covariance of review biases
    b = pd.Series(np.dot(K_b, alpha), index=X2.index)
    covb = alpha_f*(K_b - np.dot(K_b, np.dot(Kinv, K_b)))

    reviewer_bias = pd.Series(np.dot(np.diag(1./X2.sum(0)), np.dot(X2.T, b)), index=X2.columns, name='ReviewerBiasMean')
    reviewer_bias_std = pd.Series(np.dot(np.diag(1./X2.sum(0)), np.dot(X2.T, np.sqrt(np.diag(covb)))), index=X2.columns, name='ReviewerBiasStd')

    fig, ax = plt.subplots()
    reviewer_bias.hist(bins=100, ax=ax)
    _ = ax.set_title(who + ' Calibration Adjustments Histogram', fontsize=18)

    # Sanity Check
    raw_score = pd.Series(np.dot(np.diag(1./X1.sum(0)), np.dot(X1.T, reviews[column])), index=X1.columns)
    prob_accept[prob_accept==0] = 1/(10*samples)
    prob_accept[prob_accept==1] = 1-1/(10*samples)

    fig, ax = plt.subplots()
    ax.plot(np.linspace(-8,9, len(prob_accept)), prob_accept.sort_values())
    ax.set_title('Probability of accept', fontsize=18)
    ax.set_xlabel('Raw score', fontsize=16)
    _ = ax.set_ylabel('Probability of accept', fontsize=16)

    fig, ax = plt.subplots()
    ax.plot(raw_score, np.log(prob_accept)- np.log(1-prob_accept), 'rx')
    ax.set_title('Raw Score vs. Log odds of accept', fontsize=18)
    ax.set_xlabel('Raw score', fontsize=16)
    _ = ax.set_ylabel('Log odds of accept', fontsize=16)

    fig, ax = plt.subplots()
    ax.plot(reviews[column], reviews[cal_name], 'rx')
    ax.set_xlim([-9, 9])
    ax.set_title('Original vs. Calibrated ' + who + ' scores', fontsize=18)
    ax.set_xlabel('Original ' + who + ' score', fontsize=16)
    _ = ax.set_ylabel('Calibrated ' + who + ' score', fontsize=16)


    # Apply Laplace smoothing to accept probabilities before incorporating them.
    k = 5
    S = X2.shape[1]
    num = round(1/S,4)
    den = round(1+k/S,3)
    revs = reviews.join((prob_accept+num)/den, on='ID').join(reviewer_bias, on='Email').join(reviewer_bias_std, on='Email').sort_values(by=[prob_name,'ID', cal_name], ascending=False)
    reviews = reviews.join((prob_accept+num)/den, on='ID')


    # Save the computed information to disk
    #revs.to_csv(os.path.join(cu.cmt_data_directory, date + '_processed_reviews.csv'), encoding='utf-8')

    return reviews, revs



# %% #### FILES FUNCTIONS ####

# Function for getting the Paper ID number from file name
def getID(name):
    """
    Function used to get the paper ID from its file name.

    Parameters
    ----------
    name : string
        File name of the form: "Paper_xxx.pdf" or "Submission_xxx.pdf", etc.

    Returns
    -------
    Integer
        Paper ID.

    """
    num = name.split('_')[1].split('.')[0]

    return int(num)



# Function used to count the number of pages in the pdf file
def countPage(path_src, files, df):
    """
    Function used to count the number of pages from pdf files and check if they
    are compliant to the conference rules.

    Parameters
    ----------
    path_src : string
        Path of the main paper folders.
    files : string
        List of the folders.
    df : DataFrame
        DataFrame where to pack all the extracted information.

    Returns
    -------
    df : DataFrame
        DataFrame with all the extracted information.

    """
    for file in files:
        file_name = path_src + file
        doc = pdf.open(file_name)
        np = doc.page_count

        paper_ID = getID(file)
        if (np > 8):
            over_len = 'Yes'
        else:
            over_len = 'No'

        if (np > 10):
            compl = 'No'
        else:
            compl = 'Yes'

        new_row = {'ID':paper_ID, 'Pages':np, 'Overlength':over_len, 'Compliant':compl}

        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)

    return df



# %% #### SUBMISSION FUNCTIONS ####

# Function to get the list of authors from a submission
def getAuthorList(authors):
    """
    Function used to extract the list of all the authors of a paper.

    Parameters
    ----------
    authors : Series
        Authors information extracted from CMT.

    Returns
    -------
    new_author_list : List
        List of all authors for a paper.

    """
    new_author = authors.replace('*','')
    author_list = new_author.split(';')

    new_author_list = [a+',' for a in author_list[:-1]] + [author_list[-1]] # Last modification to insert a comma

    return new_author_list



# Function to get the list of authors from a submission (modified version)
def getAuthorList2(authors):
    """
    Function used to extract the list of all the authors of a paper.

    Parameters
    ----------
    authors : Series
        Authors information extracted from CMT.

    Returns
    -------
    author_list : List
        List of all authors for a paper.

    """
    new_author = authors.replace('*','')
    author_list = new_author.split(';')

    return author_list




# Function to get both Author and Affiliation lists
def getAuthorOrganizationList(authors):
    """
    Function used to extract both the author and organization information for a paper.

    Parameters
    ----------
    authors : Series
        Authors information extracted from CMT.

    Returns
    -------
    aut : List
        List of all the authors for a paper.
    org : List
        List of all the organizations for a paper.

    """
    authors_list = getAuthorList(authors)
    aut = []
    org = []
    for a in authors_list:
        aut.append(a.split('(')[0][:-1].strip())
        org.append(a.split('(')[-1][:-2].strip())

    org[-1] = org[-1] + authors_list[-1][-2]

    return aut, org



# Function to get the list of author emails from a submission
def getEmailList(emails):
    """
    Function used to extract all the authors' emails for a paper.

    Parameters
    ----------
    emails : Series
        Email information extracted from CMT.

    Returns
    -------
    email_list : List
        List of all the emails for a paper.

    """
    new_email = emails.replace('*','')
    email_list = new_email.split(';')

    return email_list



# Function to get subject areas from a submission
def getSubjects(subject):
    """
    Function to get subject areas of the form "Field -> Area" from a submission.

    Parameters
    ----------
    subject : String
        Subject Area of the form "Field -> Area", i.e. "Parent subject -> Child subject".

    Returns
    -------
    s1 : String
        Parent subject.
    s2 : String
        Child subject.

    """
    sub = subject.split('->')
    s1 = sub[0][:-1]  # Parent subject
    s2 = sub[1][1:]   # Child subject

    return s1, s2



# Function to get the affiliation of the principal author
def getMainOrganization(authors):
    """
    Function used to extract the organization of the principal author.

    Parameters
    ----------
    authors : String
        String of the list of all the authors' organizations.

    Returns
    -------
    org : String
        Organization of the principal author.

    """
    org = authors.split('*')[0].split(';')[-1].split('(')[-1][:-1]

    return org



# Function to retrieve the country of each submission
def count_country(users, papers, country_names, sort='value'):
    """
    Function used for retrieving the country for each submission.

    Parameters
    ----------
    users : DataFrame
        DataFrame containing all the User information.
    papers : DataFrame
        DataFrame containing all the Paper information.
    country_names : List
        List of contributing countries.
    sort : String, optional
        String for ordering the result:
            - 'value' : the result is order by the number of received papers in descending order
            - 'index' : the result is ordered by th ecountry name in alphabetical order
        The default is 'value'.

    Returns
    -------
    paper_per_country : Series
        A Series containing the (ordered) number of papers per contry.

    """
    paper_country = {el:0 for el in country_names}

    i = 0
    while i < len(papers):
        emails = getEmailList(papers.iloc[i]['Author Emails'])
        for em in emails:
            try:
                if str(users.loc[em.strip()]['Country']) != 'nan':
                    count = users.loc[em.strip()]['Country']
                    paper_country[count] += 1
                    break
                else:
                    continue
            except KeyError:
                continue
        i += 1

    if sort == 'value':
        paper_per_country = pd.Series(paper_country).sort_values(ascending=False)
    elif sort == 'index':
        paper_per_country = pd.Series(paper_country).sort_index()
    else:
        print('Error: sort index unsupported!')

    return paper_per_country




# %% #### PROGRAM FUNCTIONS ####

from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


# Function for writing the current session information in a box
def make_session_header(document, time1, time2, session, room, chair):
    """
    Function used to create the header for each session in the program. Specifically,
    the session header includes in a gray box, the following information:
        ___________________________________________________
        Start time -- End time
        Session name
        Room: room name
        Session Chair(s): Chair name
        ___________________________________________________

    Parameters
    ----------
    document : Document
        Hanlde of the document where to write the session information.
    time1 : String
        Start time of the session.
    time2 : String
        End time of the session.
    session : String
        Session name.
    room : String
        Room of the session.
    chair : String
        Namae of the session Chair.

    Returns
    -------
    None.

    """
    table = document.add_table(rows=1, cols=1)
    hdr_cells = table.rows[0].cells
    ch = None if chair == '' else chair
    if ch:
        header = time1 + '--' + time2 + '\n' + session + '\n' + 'Room: ' + room + '\n' + 'Session Chair(s): ' + ch
    else:
        header = time1 + '--' + time2 + '\n' + session + '\n' + 'Room: ' + room
    # header = str(time1) + '--' + str(time2) + '\n' + session + '\n' + 'Room: ' + room + '\n' + 'Session Chair: ' + chair
    hdr_cells[0].text = header
    table.style = 'Medium Grid 1'



# Function to list all papers of a session in the program
def make_session(prog, sessions, papers, document):
    """
    Function used to create the program of a single session. Specifically, the
    function writes a list of all the paper to be presented in the session, using
    the following format:
        Start time  Paper title
        Author 1 (Organization 1), Author 2 (Organization 2), Author 3 (Organization 3), ...

    Parameters
    ----------
    prog : DataFrame
        DataFrame containing the program information for each session. This function
        is called within a groupby() function, hence containing a single session.
    sessions : DataFrame
        DataFrame containing all the session information.
    papers : DataFrame
        DataFrame containing all the papers information.
    document : Document
        Hanlde of the document where to write the session information.

    Returns
    -------
    None.

    """
    p = document.add_paragraph('')
    p.add_run('').italic = True

    # Writing the Session information
    s = prog.iloc[0]['Session ID']
    make_session_header(document, sessions.loc[s]['Start Time'].strftime('%H:%M'),
                        sessions.loc[s]['End Time'].strftime('%H:%M'),
                        sessions.loc[s]['Session Name'],
                        sessions.loc[s]['Room'],
                        chair=sessions.loc[s]['Chair'])
    p = document.add_paragraph('')
    p = document.add_paragraph('')

    # Write the paper information inside the session
    for i in range(len(prog)):

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        title = p.add_run(str(prog.iloc[i]['Start Time'].strftime('%H:%M')) + ' ' + papers.loc[int(prog.iloc[i]['Paper ID'])]['Title'])
        title.bold = True
        title.italic = True
        p.paragraph_format.space_after = 1
        p = document.add_paragraph('')
        p.paragraph_format.space_after = 1
        p.add_run(getAuthorList(papers.loc[int(prog.iloc[i]['Paper ID'])]['Authors']))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        p = document.add_paragraph('')



# Function to print a timetable for each session
def print_prog(prog):
    """
    Function used for printing a time-table for each session (for DEBUG purpose).

    Parameters
    ----------
    prog : DataFrame
        DataFrame containing all the program information (Paper ID, Start time, and End time).

    Returns
    -------
    None.

    """
    print('Session ID:', prog.iloc[0]['Session ID'])
    for i in range(len(prog)):
        print(prog.iloc[i]['Paper ID'], '\b:', prog.iloc[i]['Start Time'].strftime('%H:%M'),
              '--', prog.iloc[i]['End Time'].strftime('%H:%M'))
        # s = f"{prog.iloc[i]['Paper ID']:{4}.{0}f}: {prog.iloc[i]['Start Time'].strftime('%H:%M')} -- {prog.iloc[i]['End Time'].strftime('%H:%M')}"
        # print(s)
    print(' ')



# Function to make a list of all video presentations
def make_video_list(videos, papers, document):
    """
    Funtion used to generate a document with the list of all papers which will
    be presented remotely by a video, using the following format:
        Paper title
        Author 1 (Organization 1), Author 2 (Organization 2), Author 3 (Organization 3), ...

    Parameters
    ----------
    videos : DataFrame
        DataFrame containing all paper IDs presented remotely.
    papers : DataFrame
        DataFrame containing all the papers information.
    document : Document
        Hanlde of the document where to write the session information.

    Returns
    -------
    None.

    """
    p = document.add_paragraph('')
    p.add_run('').italic = True

    for i in range(len(videos)):
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            title = p.add_run(papers.loc[videos[i]]['Title'])
            title.bold = True
            title.italic = True
            p.paragraph_format.space_after = 1
            p = document.add_paragraph('')
            p.paragraph_format.space_after = 1
            p.add_run(getAuthorList(papers.loc[videos[i]]['Authors']))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(10)
            p = document.add_paragraph('')
        except KeyError:
            continue
