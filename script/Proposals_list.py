# -*- coding: utf-8 -*-
"""
This script is used for manipulating the Special Session/Tutorial/Competition/Workshop
proposals received to IJCNN 2025 conference. Specifically, it makes three tasks:
    1. First, the script moves files from a tree-based source folder to a single
       folder and rename them according to the Paper ID.
    2. Second, it extract uself information from the CMT report exporting a
       simple xlsx file.
    3. Third, a list in .docx format of all submitted proposals is generated.

It is used for moving files downloaded from Microsoft CMT system.

Created on Thu Nov 28 17:34:35 2024

@author: Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
"""


import os
import shutil
import pandas as pd
import IJCNN_utils as ut
from docx import Document


# %% 1. Move and rename proposals

# Settings
path_src = './data/Submissions/'
path_dst = './data/Proposals/'

# Sub-folder
sub_type = '/Submission/'
# sub_type = '/CameraReady/'
# sub_type = '/Copyright/'

copy = 0  # 0: move files,  1: copy files

# Scan the current folder
folders  = os.listdir(path_src)
N_papers = len(folders)

# Loop for moving files
for i in range(len(folders)):
    p1 = path_src + folders[i] + sub_type
    name_src = os.listdir(p1)[0]
    ext = '.' + name_src.split('.')[-1]
    name_dst = 'Proposal_' + folders[i] + ext
    if copy:
        shutil.copy2(p1+name_src, path_dst+name_dst)
    else:
        shutil.move(p1+name_src, path_dst+name_dst)
        os.rmdir(p1)

print("End of moving {} files".format(N_papers))



# %% 2. Extract information and generate a simple report

# Settings
path_src = './data/Proposals/'
path_dst = './report/'

# Proposal type
proposal = 'Special Session'
# proposal = 'Workshop'
# proposal = 'Tutorial'
# proposal = 'Competition'

input_file = path_src + proposal + '_Proposals.xlsx'
save_file  = path_dst + proposal + '_Report.xlsx'


# Reading the input file
df = pd.read_excel(input_file)
N_papers = df.shape[0]

paperID = df['Paper ID']
titles  = df['Paper Title']
authors = df['Authors']


# Create the list of proposals
paper_list = pd.DataFrame(columns = ['ID', 'Authors', 'Title'])

author_list = []
for author in authors:
    paper_author = ut.getAuthorList(author)
    author_list.append(paper_author)

for n, paper in enumerate(author_list):
    nome = ''
    for auth in paper:
        nome = nome + auth
    new_row = {'ID':paperID[n], 'Authors':nome, 'Title':titles[n]}
    paper_list = pd.concat([paper_list, pd.DataFrame([new_row])], ignore_index=True)


# Save Excel report
paper_list.to_excel(save_file, index=False)
print("End of exporting {} papers".format(N_papers))



# %% 3. Generate a list in .docx format of all submitted proposals

document = Document()

# Write the document header
head_title = 'IJCNN 2025 ' + proposal + ' Proposals'
document.add_heading(head_title, 0)

# Write each proposal
for i in range(N_papers):
    p = document.add_paragraph('')
    p.add_run(str(i+1) + '. Paper ID: ' + str(paper_list.iloc[i]['ID'])).bold = True
    p = document.add_paragraph('')
    p.add_run(paper_list.iloc[i]['Title']).bold = True
    p = document.add_paragraph('')
    p.add_run(paper_list.iloc[i]['Authors']).italic = True
    p = document.add_paragraph('')


# Save the docx document
doc_title = path_dst + proposal + '_List.docx'
document.save(doc_title)
