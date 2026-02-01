# -*- coding: utf-8 -*-
"""
This script is used to generate the conference program from the report of accepted
papers containing the full information.
It is used with three excel files:
- a file downloaded from Microsoft CMT system containing all the accepted papers,
- a file containing sessions' information organized with the following fields:
     Session ID - Session Name - Room - Start Time - End Time - Day - Chair
- a file with the program content organized with the following fields:
     Session ID - Paper ID - Start Time - End Time

Created on Mon May 19 20:14:07 2025

@author: Michele Scarpiniti -- DIET Dpt. (Sapienza University of Rome)
"""

import pandas as pd
import IJCNN_utils as ut
from docx import Document


# Conference days
days = ['June 30', 'July 01', 'July 02', 'July 03', 'July 04', 'July 05']


# Set file names
path = './data/'
dest = './report/'
paper_file = path + 'All_Accepted_Papers.xlsx'
session_file = path + 'Sessions.xlsx'
program_file = path + 'Program.xlsx'

save_file = dest + 'Program.docx'

# Create a new document
composed = Document()

# Load paper data
papers = ut.load_submissions(paper_file)
papers.set_index('ID', inplace=True)


# Loop over all days
for day in days:
    save_file_day  = dest + 'Program_' + day + '.docx'

    # Load all data
    sessions = pd.read_excel(session_file, sheet_name=day)
    program  = pd.read_excel(program_file, sheet_name=day)

    # Set the new index
    sessions.set_index('Session ID', inplace=True)

    # Create a new document
    document = Document()

    # Add the document header
    document.add_heading('IJCNN 2025 Program -- ' + day, 0)

    p = document.add_paragraph('')
    p.add_run('').italic = True

    # Make the program
    program.groupby(by='Session ID').apply(lambda x: ut.make_session(x, sessions, papers, document))

    # Save the docx document
    document.save(save_file_day)

    print('Done for day:', day)

    # Append the current document to the full program
    for element in document.element.body:
            composed.element.body.append(element)


# Save the full program docx document
composed.save(save_file)
